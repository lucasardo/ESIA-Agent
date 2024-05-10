import os
import streamlit as st
from azure.search.documents import SearchClient
from azure.search.documents.models import VectorizedQuery
from azure.core.credentials import AzureKeyCredential
from llama_index.embeddings.azure_openai import AzureOpenAIEmbedding
from llama_index.core.schema import TextNode
from llama_index.core import VectorStoreIndex, PromptHelper, ServiceContext
from llama_index.llms.azure_openai import AzureOpenAI
from llama_index.core.settings import Settings
from docx import Document as Docx
from docx.shared import RGBColor

from st_helper import *

################################################################################################################

# Configure the baseline configuration of the OpenAI library for Azure OpenAI Service.
azure_endpoint = st.secrets["azure_endpoint"]
openai_api_key = st.secrets['openai_api_key']

openai_deployment_name = st.secrets['openai_deployment_name']
openai_api_version = st.secrets['openai_api_version']
embedding_model = st.secrets['embedding_model']
embedding_deployment_name = st.secrets['embedding_deployment_name']

search_endpoint = st.secrets['search_endpoint']
search_api_key = st.secrets['search_api_key']
search_api_version = st.secrets['search_api_version']
search_service_name = st.secrets['search_service_name']

search_url = f"https://{search_service_name}.search.windows.net/"
search_credential = AzureKeyCredential(search_api_key)

index_name = "esias-base-index"
max_tokens = 4096
dimensionality = 1536

embed_model = AzureOpenAIEmbedding(
    model=embedding_model,
    deployment_name=embedding_deployment_name,
    api_key=openai_api_key,
    azure_endpoint=azure_endpoint,
    api_version=openai_api_version
)

llmChat = AzureOpenAI(
    engine=str(openai_deployment_name),
    azure_endpoint=str(azure_endpoint),
    api_key=str(openai_api_key),
    api_version=str(openai_api_version)
)

prompt_helper = PromptHelper(
    context_window=1536,
    num_output=256,
    chunk_overlap_ratio=0.1,
    chunk_size_limit=None,
)

service_context = ServiceContext.from_defaults(
            llm=llmChat,
            embed_model=embed_model,
            prompt_helper=prompt_helper,
)

Settings.llm = llmChat
Settings.embed_model = embed_model

###############################################################################

# Set web page title, icon, and layout
st.set_page_config(
    page_title="ESIA Agent 💬 WSP",
    page_icon=":robot:",
    layout="wide"  # Set layout to wide for better organization
)

# Sidebar
st.sidebar.image("https://download.logo.wine/logo/WSP_Global/WSP_Global-Logo.wine.png", width=100)

# Button to start a new thread (restart the app)
if st.sidebar.button("New Thread"):
    st.experimental_rerun()
    user_input = ""

# History of user inputs
st.sidebar.markdown('#') 
st.sidebar.header("History")
st.sidebar.write("- I have to write an ESIA for an offshore wind farm project, called 'ItaWind'. The wind farm will be located in Italy, off the coast of Molise.")

st.write("<h1 style='color: #F9423A; text-align: center;'>Hi, how can I help you today?</h1>", unsafe_allow_html=True)
st.write("<h5 style='color: #F9423A; text-align: center;'>I am your assistant for drafting ESIA reports. Give me information on the name of the project, type of infrastructure, location, and I will lay out a draft for you.</h5>", unsafe_allow_html=True)

st.markdown('#') 

######################################
####### PROJECT MANAGER
######################################

filter = ""

st.markdown('#') 
st.markdown('#') 

init_prompt = st.text_input("Ask anything")

# React to user input
if init_prompt:
    
    st.write("Request for Proposal:", init_prompt)
    
    indexes = [index_name]

    question = str(init_prompt)
    
    prompt = AGENT_INTRO_PROMPT
    
    tools = [GetDocSearchResults_Tool(
    indexes=indexes, k=10, reranker_th=1, sas_token='na')]
    
    COMPLETION_TOKENS = 4096

    llm = AzureChatOpenAI(deployment_name=openai_deployment_name, openai_api_version=openai_api_version,
                            openai_api_key=openai_api_key, azure_endpoint=azure_endpoint, temperature = 0)
   
    agent = create_openai_tools_agent(llm, tools, prompt)

    agent_executor = AgentExecutor(
    agent=agent, tools=tools, handle_parsing_errors=True, verbose=False)

    with_message_history = RunnableWithMessageHistory(
    agent_executor,
    get_session_history,
    input_messages_key="question",
    history_messages_key="history"
    )
    
    #Prova
    session_id = 125

    response = with_message_history.invoke(
        {"question": question},
        config={"configurable": {"session_id": session_id}}
    )

    st.write("<h2 style='color: #F9423A;'>INTRODUCTION", unsafe_allow_html=True)

    history = update_history(session_id, question, response["output"], indexes)

    full_response = {
        "question": question,
        "output": response["output"],
        "history": history
    }

    response_text = full_response['output']
    response_intro = f"{response_text}"
    
    st.markdown(response_intro)
    search_results = simple_hybrid_search(question, index_name, filter, search_url, search_credential, azure_endpoint, openai_api_key, openai_api_version, embedding_deployment_name)

    # Initialize an empty list to store the nodes
    sources_nodes = []
    for result in search_results:
        score = result["@search.score"]
        path = result["doc_path"]
        # Create a dictionary for the node
        node = {"path": path, "score": score}
        # Append the dictionary to the list of nodes
        sources_nodes.append(node)
    
    with st.expander("See sources"):
        for node in sources_nodes[0:5]:
            file_name = os.path.basename(node["path"])
            st.write("Source file: ", file_name)
            raw_score = node["score"]
            score = "{:.2f}".format(raw_score)
            score = float(score)
            if score >= 0.02:
                st.write(":heavy_check_mark: High confidence. Search score: ", score)
            elif score >= 0.01:
                st.write(":warning: Medium confidence. Search score: ", score)
            else:
                st.write(":x: Low confidence. Search score: ", score)
    
######################################
####### ENVIRONMENTAL ENGINEER
######################################
                
    st.markdown('#')
    
    st.write("<h2 style='color: #F9423A;'>ENVIRONMENTAL IMPACT", unsafe_allow_html=True)
    
    question = str(init_prompt)
    
    prompt = AGENT_ENV_PROMPT
    
    tools = [GetDocSearchResults_Tool(
    indexes=indexes, k=10, reranker_th=1, sas_token='na')]
    
    COMPLETION_TOKENS = 4096

    llm = AzureChatOpenAI(deployment_name=openai_deployment_name, openai_api_version=openai_api_version,
                            openai_api_key=openai_api_key, azure_endpoint=azure_endpoint, temperature = 0)

    agent = create_openai_tools_agent(llm, tools, prompt)

    agent_executor = AgentExecutor(
    agent=agent, tools=tools, handle_parsing_errors=True, verbose=False)

    with_message_history = RunnableWithMessageHistory(
    agent_executor,
    get_session_history,
    input_messages_key="question",
    history_messages_key="history"
    )
    
    #Prova
    session_id = 126

    response = with_message_history.invoke(
        {"question": question},
        config={"configurable": {"session_id": session_id}}
    )

    history = update_history(session_id, question, response["output"], indexes)

    full_response = {
        "question": question,
        "output": response["output"],
        "history": history
    }

    response_text = full_response['output']
    response_content = f"{response_text}"
    
    st.markdown(response_content)
    search_results = simple_hybrid_search(question, index_name, filter, search_url, search_credential, azure_endpoint, openai_api_key, openai_api_version, embedding_deployment_name)

    # Initialize an empty list to store the nodes
    sources_nodes = []
    for result in search_results:
        score = result["@search.score"]
        path = result["doc_path"]
        # Create a dictionary for the node
        node = {"path": path, "score": score}
        # Append the dictionary to the list of nodes
        sources_nodes.append(node)
    
    with st.expander("See sources"):
        for node in sources_nodes[0:5]:
            file_name = os.path.basename(node["path"])
            st.write("Source file: ", file_name)
            raw_score = node["score"]
            score = "{:.2f}".format(raw_score)
            score = float(score)
            if score >= 0.02:
                st.write(":heavy_check_mark: High confidence. Search score: ", score)
            elif score >= 0.01:
                st.write(":warning: Medium confidence. Search score: ", score)
            else:
                st.write(":x: Low confidence. Search score: ", score)

######################################
####### ENVIRONMENTAL ECONOMIST
######################################
                
    st.markdown('#')
    
    st.write("<h2 style='color: #F9423A;'>SOCIAL IMPACT", unsafe_allow_html=True)
    
    question = str(init_prompt)
    
    prompt = AGENT_SOCIAL_PROMPT
    
    tools = [GetDocSearchResults_Tool(
    indexes=indexes, k=10, reranker_th=1, sas_token='na')]
    
    COMPLETION_TOKENS = 4096

    llm = AzureChatOpenAI(deployment_name=openai_deployment_name, openai_api_version=openai_api_version,
                            openai_api_key=openai_api_key, azure_endpoint=azure_endpoint, temperature = 0)

    agent = create_openai_tools_agent(llm, tools, prompt)

    agent_executor = AgentExecutor(
    agent=agent, tools=tools, handle_parsing_errors=True, verbose=False)

    with_message_history = RunnableWithMessageHistory(
    agent_executor,
    get_session_history,
    input_messages_key="question",
    history_messages_key="history"
    )
    
    session_id = 127

    response = with_message_history.invoke(
        {"question": question},
        config={"configurable": {"session_id": session_id}}
    )

    history = update_history(session_id, question, response["output"], indexes)

    full_response = {
        "question": question,
        "output": response["output"],
        "history": history
    }

    response_text = full_response['output']
    response_social = f"{response_text}"
    
    st.markdown(response_social)
    search_results = simple_hybrid_search(question, index_name, filter, search_url, search_credential, azure_endpoint, openai_api_key, openai_api_version, embedding_deployment_name)

    # Initialize an empty list to store the nodes
    sources_nodes = []
    for result in search_results:
        score = result["@search.score"]
        path = result["doc_path"]
        # Create a dictionary for the node
        node = {"path": path, "score": score}
        # Append the dictionary to the list of nodes
        sources_nodes.append(node)
    
    with st.expander("See sources"):
        for node in sources_nodes[0:5]:
            file_name = os.path.basename(node["path"])
            st.write("Source file: ", file_name)
            raw_score = node["score"]
            score = "{:.2f}".format(raw_score)
            score = float(score)
            if score >= 0.02:
                st.write(":heavy_check_mark: High confidence. Search score: ", score)
            elif score >= 0.01:
                st.write(":warning: Medium confidence. Search score: ", score)
            else:
                st.write(":x: Low confidence. Search score: ", score)

######################################
####### SUMMARIZER
######################################
                
    st.markdown('#')
    
    st.write("<h2 style='color: #F9423A;'>CONCLUSION", unsafe_allow_html=True)
    
    question = str(init_prompt)
    
    prompt = AGENT_CONCLUSION_PROMPT
    
    tools = [GetDocSearchResults_Tool(
    indexes=indexes, k=10, reranker_th=1, sas_token='na')]
    
    COMPLETION_TOKENS = 4096

    llm = AzureChatOpenAI(deployment_name=openai_deployment_name, openai_api_version=openai_api_version,
                            openai_api_key=openai_api_key, azure_endpoint=azure_endpoint, temperature = 0)

    agent = create_openai_tools_agent(llm, tools, prompt)

    agent_executor = AgentExecutor(
    agent=agent, tools=tools, handle_parsing_errors=True, verbose=False)

    with_message_history = RunnableWithMessageHistory(
    agent_executor,
    get_session_history,
    input_messages_key="question",
    history_messages_key="history"
    )
    
    session_id = 128

    response = with_message_history.invoke(
        {"question": question},
        config={"configurable": {"session_id": session_id}}
    )

    history = update_history(session_id, question, response["output"], indexes)

    full_response = {
        "question": question,
        "output": response["output"],
        "history": history
    }

    response_text = full_response['output']
    response_conclusion = f"{response_text}"
    
    st.markdown(response_conclusion)

############################################################################

#Save as word

    def save_as_word(response_intro, response_content, response_social, response_conclusion):
        # Create a new Word document
        doc = Docx()
        
        # Function to add content with headings
        def add_content_with_headings(content):
            lines = content.split('\n')
            for line in lines:
                if line.startswith('####'):
                    # Add heading
                    heading = doc.add_heading(line.lstrip('# '), level=2)
                    # Set heading color to red
                    for run in heading.runs:
                        run.font.color.rgb = RGBColor(255, 0, 0)  # Red color
                elif line.startswith('###'):
                    # Add heading
                    # Add heading
                    heading = doc.add_heading(line.lstrip('# '), level=1)
                    # Set heading color to red
                    for run in heading.runs:
                        run.font.color.rgb = RGBColor(255, 0, 0)  # Red color
                else:
                    # Add regular paragraph
                    doc.add_paragraph(line)
            # Add empty paragraph for separation
            doc.add_paragraph()
    
        # Add markdown to the document with headings
        add_content_with_headings(response_intro)
        add_content_with_headings(response_content)
        add_content_with_headings(response_social)
        add_content_with_headings(response_conclusion)
    
        # Save the document
        doc.save("ESIA Draft.docx")
    
    save_as_word(response_intro, response_content, response_social, response_conclusion)
    with open("ESIA Draft.docx", "rb") as f:
        bytes_data = f.read()
    st.download_button(
        label="Click here to download",
        data=bytes_data,
        file_name="ESIA Draft.docx",
        mime="application/docx"
    )

